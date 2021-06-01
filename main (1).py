import urllib.request
import cv2
import matplotlib.pyplot as plt
import smtplib
import sys
import requests
import json
import re
from bs4 import BeautifulSoup
from PIL import Image
import docx

booksource= 'https://www.gutenberg.org/files/1342/1342-0.txt'
bookimagesource= 'https://upload.wikimedia.org/wikipedia/commons/3/37/Pride_and_Pr%C3%A9judice_CH_43.jpg'
imagepath = "image.jpg"

def task7(avg=0):
    # Write to docx
    document = docx.Document()
    document.add_heading('Title of the book: Pride and Prejudice', 4)
    document.add_heading('Author of the book: Jane Austen', 4)
    document.add_heading('Report by Wahee', 5)
    document.add_picture(r'C:\Users\Deniz\Desktop\Deniz\SL11\result.jpg',
                         width=docx.shared.Inches(2.5),
                         height=docx.shared.Inches(2.5))
    document.add_page_break()
    document.add_picture(r'C:\Users\Deniz\Desktop\Deniz\SL11\plot.png',
                         width=docx.shared.Inches(2.5),
                         height=docx.shared.Inches(2.5))
    document.add_paragraph('The plot demonstrates the count of paragraphs with given count of words.')
    document.add_paragraph(f'Avg number of words in the paragraphs of first chapter: {avg}')
    document.add_paragraph('This is carried out only for the first chapter of the book Pride and Prejudice by Jane Austen')
    document.add_paragraph(f'Book source: {booksource}')
    document.add_paragraph(f'Book image source: {bookimagesource}')

    document.save('new.docx')


def getBookInfo():
    # Parse html to get book information
    response = requests.get(booksource)
    response.encoding = response.apparent_encoding
    book = {}
    title = re.findall(r'(?<=Title: ).+(?=\r)', response.text)[0]    # Regex to find between '>'  and  '<'
    author = re.findall(r'(?<=Author: ).+(?=\r)', response.text)[0]    # Regex to find between '>'  and  '<'
    firstChapter = re.findall('(?<=Chapter 1\r).+?(?=Chapter 2)', response.text, flags=re.S)
    book['title'] = title
    book['author'] = author
    book['first chapter'] = firstChapter[1]

    # Returns book informaiton
    return book


def plotParagraphs(book):
    # Create the plot
    paragraphs = book['first chapter'].split('\n')
    wordCounts = []
    length = 0
    max = 0
    for paragraph in paragraphs:
        print(paragraph)
        length = len(paragraph)
        if length < 2:
            continue
        elif length > max:
            max = length
        length -= length % 10
        wordCounts.append(length)
    avg = 0
    y = [0 for i in range(int(max/10)+1)]
    x = [i for i in range(0, 10*(int(max / 10) + 1),10)]
    for c in wordCounts:
        y[int(c/10)] += 1
        avg += c
    avg /= len(wordCounts)

    # Prepare the plot
    plt.plot(x,y)
    plt.xlabel('Word count')
    plt.ylabel('Para count')
    plt.title('Frequency of paragraphs with given length')
    plt.savefig('plot.png')     # Save it as an image
    plt.show()
    return avg                  # Returns the avg to be written in the docx file


def download_web_image(url, saveto="image.jpg"):
    urllib.request.urlretrieve(url, saveto)


def task5_6():
    # Get the images
    download_web_image(bookimagesource)
    path = "image.jpg"
    img = cv2.imread(imagepath)
    img2 = cv2.imread('im2.jpg')

    # Resize crop, sum images
    img2_resized = cv2.resize(img2, (450,450), interpolation=cv2.INTER_AREA)
    imgCropped = img[300:750, 0:450]
    last = cv2.addWeighted(imgCropped, 0.5, img2_resized, 0.5, 0)

    # Text styling
    font = cv2.FONT_HERSHEY_SIMPLEX
    fontScale = 1
    color = (255, 0, 0)
    thickness = 2

    # Putting in the texts
    last = cv2.putText(last, 'Pride and Prejudice', (70, 200), font, fontScale, color, thickness, cv2.LINE_AA)
    last = cv2.putText(last, 'Jane Austen', (140, 400), font, fontScale, color, thickness, cv2.LINE_AA)

    # Demonstration of images
    # imgRotated = cv2.rotate(imgCropped, cv2.ROTATE_90_CLOCKWISE)
    # cv2.imshow("rotated", imgRotated)
    # cv2.imshow("last", last)
    # cv2.waitKey(0)

    # Save the result
    cv2.imwrite('result.jpg', last)



if __name__ == '__main__':
    book = getBookInfo()
    avg = plotParagraphs(book)
    task5_6()
    task7(avg=avg)

